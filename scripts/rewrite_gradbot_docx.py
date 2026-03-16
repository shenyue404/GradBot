from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(r"d:\AAA大创\GradBot")
OUTPUT = ROOT / "灵犀毕设-GradBot-更新版.docx"


def add_title(doc: Document, text: str, size: int, bold: bool = True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_heading(doc: Document, text: str, level: int):
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str):
    p = doc.add_paragraph(text)
    fmt = p.paragraph_format
    fmt.first_line_indent = Pt(24)
    fmt.line_spacing = 1.5
    return p


def add_list(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def build_document():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)

    add_title(doc, "毕设灵犀GradBot系统", 18)
    add_title(doc, "软件使用说明书", 16)
    add_title(doc, "[ V1.0（更新版） ]", 12, bold=False)
    add_title(doc, "[ 著作人：李欣艳 ]", 12, bold=False)
    add_title(doc, "[ 2026.03 ]", 12, bold=False)

    add_heading(doc, "1 引言", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(
        doc,
        "当前高校毕业设计管理过程中，师生在选题申报、任务书编写、开题报告撰写、中期检查、教师审核与过程跟踪等环节普遍面临流程长、协同复杂、反馈不及时、质量标准不统一等问题。传统做法对教师经验依赖较强，且难以为学生提供连续、细粒度的指导支持。"
    )
    add_para(
        doc,
        "本说明书用于介绍“毕设灵犀GradBot——基于大模型的高校毕业设计全过程智能管理系统”的系统用途、运行环境、主要功能及使用方式，帮助学生、教师与管理员快速了解并正确使用系统。当前版本已经实现学生端、教师端、管理员端三类角色的核心流程，能够支撑毕业设计过程中的智能生成、审核反馈、人员管理与进度预警。"
    )

    add_heading(doc, "1.2 背景", 2)
    add_para(
        doc,
        "毕业设计是高校本科人才培养中的关键环节，涉及选题、任务书、开题、中期检查、评审与答辩等多个阶段，数据流转频繁，材料编写工作量大。现有通用教学管理系统多以流程记录为主，缺乏针对毕业设计全过程的深度智能辅助能力。"
    )
    add_para(
        doc,
        "“毕设灵犀GradBot”系统结合 MySQL 数据管理能力、Web 前后端分离架构以及 DeepSeek 大模型接口，构建了覆盖学生生成、教师审核、管理员调配和进度监测的一体化平台。系统重点解决毕业设计过程中的文档生成效率低、教师审核压力大、师生关系难追踪、风险学生识别不及时等问题。"
    )

    add_heading(doc, "2 系统概述", 1)
    add_heading(doc, "2.1 软件用途", 2)
    add_para(
        doc,
        "本系统是一套面向高校毕业设计管理场景的智能化 Web 平台，主要服务于学生、教师和管理员三类用户。系统围绕毕业设计全过程，提供智能选题、任务书生成、开题报告生成、中期报告生成、教师 AI 评分与评语、管理员师生分配与预警分析等功能。"
    )

    add_heading(doc, "2.2 功能概述", 2)
    add_list(
        doc,
        [
            "智能选题推荐：学生输入兴趣方向与关键词后，系统自动生成 5 个符合本科生水平的中文备选题目。",
            "任务书智能生成：基于已选题目生成结构化任务书，覆盖主要任务、目的、主要内容、基本要求、前期基础、预期成果、参考文献和工作进度安排。",
            "开题报告智能生成：按照高校常见模板生成选题背景、选题目的、选题意义、国内外研究现状、主要内容和研究方法、前期基础与预期成果、论文进度安排和参考文献。",
            "中期报告智能生成：结合当前课题自动生成进度概览、已完成工作、问题与解决方案、下一阶段计划等内容。",
            "教师 AI 审核：教师可对任务书、开题报告和中期报告执行 AI 评分，并在 AI 生成评语基础上进行人工修改后提交。",
            "管理员管理功能：管理员可录入教师和学生、调整指导关系、查看教师名下学生详情，并按学院、专业、年级筛选进度与质量预警学生。",
            "AI 连通检测：系统提供 AI 连通性检测页面，便于快速排查大模型接口状态。"
        ],
    )

    add_heading(doc, "2.3 运行环境", 2)
    add_para(doc, "操作系统：Windows 10 及以上版本。")
    add_para(doc, "后端环境：Node.js、Express、Sequelize、MySQL。")
    add_para(doc, "前端环境：React、TypeScript、Vite、Ant Design。")
    add_para(doc, "数据库：MySQL 8.0 推荐。")
    add_para(doc, "AI 服务：DeepSeek 兼容接口，需配置有效 API Key。")

    add_heading(doc, "3 功能说明", 1)
    add_heading(doc, "3.1 用户注册与登录", 2)
    add_para(
        doc,
        "用户进入系统后可在登录页面输入邮箱、密码进行登录。系统支持学生、教师和管理员三类角色，不同角色登录后将自动进入各自对应的功能界面。学生注册时可填写学号、专业等信息；教师账号和学生账号也可由管理员在后台单个录入或批量导入。"
    )

    add_heading(doc, "3.2 个人中心", 2)
    add_para(
        doc,
        "学生和教师登录后均可进入个人中心查看并更新基本信息，包括姓名、专业或职称、联系方式、研究方向等内容。系统支持修改个人资料，并在页面头部展示当前登录用户信息。"
    )

    add_heading(doc, "3.3 学生端功能", 2)
    add_para(
        doc,
        "学生端首页提供毕业设计流程状态概览，集中展示选题、任务书、开题报告和中期报告的当前状态与教师反馈。左侧导航栏包含状态总览、选题申报、任务书、开题报告、中期报告和 AI 连通状态等入口。"
    )
    add_list(
        doc,
        [
            "选题申报：学生输入兴趣领域和关键词后，系统调用 AI 生成 5 个中文选题备选项，学生可在此基础上选择并提交。",
            "任务书：学生可根据已选题目一键生成任务书初稿，并对各栏目进行编辑、保存草稿、提交审核或上传已有文档。",
            "开题报告：系统按照结构化模板生成开题报告内容，支持学生继续修改、保存和提交。",
            "中期报告：系统根据当前课题和阶段情况生成中期报告初稿，用于辅助学生梳理进展和问题。",
            "AI 连通检测：学生可进入检测页面确认后端与 DeepSeek 接口的配置状态和连通状态。"
        ],
    )

    add_heading(doc, "3.4 教师端功能", 2)
    add_para(
        doc,
        "教师端围绕审核与指导展开。系统依据管理员建立的指导关系，仅向教师展示其名下学生的待审核材料。教师左侧导航包含选题审核、任务书审核、开题审核、中期审核和 AI 连通状态等页面。"
    )
    add_list(
        doc,
        [
            "选题审核：教师可查看学生提交的选题，进行通过、退回修改或拒绝操作。",
            "任务书审核：教师可查看任务书详情，执行 AI 评分，并在 AI 评语基础上修改后提交审核意见。",
            "开题报告审核：教师可对开题报告进行 AI 评分、生成评语并提交人工审核结论。",
            "中期报告审核：教师可根据学生中期进度、问题与成果执行 AI 分析和人工审核。",
            "个人资料：教师可查看个人基础信息和指导相关资料。"
        ],
    )

    add_heading(doc, "3.5 管理员端功能", 2)
    add_para(
        doc,
        "管理员端负责全校范围内的人员管理、师生关系维护与进度质量监控。管理员登录后可进入总览页、师生分配页、学生预警页和 AI 连通状态页。"
    )
    add_list(
        doc,
        [
            "管理员总览：查看教师数、学生数、指导关系数和预警学生数等统计信息。",
            "师生分配：支持单个录入教师、学生，也支持 Excel 批量导入，并可调整教师负责的学生名单。",
            "教师名下学生详情：管理员可查看某位教师名下学生的选题、任务书、开题和中期状态，以及进度分、质量分和风险说明。",
            "进度预警：按学院、专业、年级筛选预警学生，对进度不足或质量不达标的学生进行重点标识。"
        ],
    )

    add_heading(doc, "4 系统特点", 1)
    add_list(
        doc,
        [
            "采用前后端分离架构，便于维护与扩展。",
            "支持真实 MySQL 数据库存储和多角色权限控制。",
            "接入 DeepSeek API，用于学生端智能生成与教师端 AI 审核。",
            "支持管理员建立和调整指导关系，实现教师按名下学生范围审核。",
            "支持批量导入、进度预警和 AI 连通性自检，增强系统实用性。"
        ],
    )

    add_heading(doc, "5 使用说明补充", 1)
    add_para(
        doc,
        "系统当前以 PC Web 端为主，不包含微信小程序客户端。任务书与开题报告页面采用结构化表单编辑方式，不再依赖原型中的富文本编辑形式。系统中的 AI 生成功能需要后端正确配置 DeepSeek API Key，若网络异常或接口不可达，可通过 AI 连通检测页面定位问题。"
    )
    add_para(
        doc,
        "测试环境中已预置管理员、教师和学生账号，用于流程验证和功能演示。管理员可进一步录入和调整师生数据，以适配不同学院或专业的使用场景。"
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
